
import docx
import os

doc = docx.Document()
doc.add_heading('Testing Word Support', 0)
doc.add_paragraph('This is a test paragraph from the automated test script.')
doc.add_paragraph('If you can read this, the Universal Adapter is working.')

output_path = "test_word_capability.docx"
doc.save(output_path)
print(f"Created: {os.path.abspath(output_path)}")
